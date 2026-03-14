import asyncio
import io
import json
import logging
import uuid
from datetime import datetime

from dependencies import check_document_access
from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile, status
from fastapi.responses import Response
from pydantic import BaseModel, ConfigDict
from shared.database import get_db
from shared.minio_client import download_file, generate_presigned_url, upload_file
from shared.models import AuditLog, Document, DocumentPermission, OrganizationMember, User
from shared.security import get_current_user
from sqlalchemy import func, or_, select
from sqlalchemy.ext.asyncio import AsyncSession

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/documents", tags=["documents"])

ALLOWED_EXTENSIONS = {"pdf", "docx", "xlsx", "txt"}
CONTENT_TYPES: dict[str, str] = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "txt": "text/plain",
}
MAX_FILE_SIZE = 52_428_800  # 50 MB


# ---------------------------------------------------------------------------
# Response schemas
# ---------------------------------------------------------------------------


class DocumentOut(BaseModel):
    model_config = ConfigDict(from_attributes=True)

    id: int
    title: str
    filename: str
    file_type: str
    file_size: int
    folder_path: str
    status: str
    error_message: str | None
    uploaded_by: int
    org_id: int | None
    uploaded_at: datetime
    updated_at: datetime
    indexed_at: datetime | None
    page_count: int | None
    chunk_count: int | None
    tags: list
    department: str | None


class DocumentStatusOut(BaseModel):
    id: int
    status: str
    error_message: str | None
    chunk_count: int | None


class DocumentListOut(BaseModel):
    items: list[DocumentOut]
    total: int
    page: int
    limit: int


class DocumentPreviewOut(BaseModel):
    text: str
    page_count: int | None


class PresignedUrlOut(BaseModel):
    url: str
    expires_in: int


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------


@router.post("/upload", response_model=DocumentOut, status_code=status.HTTP_201_CREATED)
async def upload_document(
    request: Request,
    file: UploadFile = File(...),
    title: str | None = Form(None),
    tags: str | None = Form(None),
    folder_path: str = Form("/"),
    org_id: int | None = Form(None),
    department: str | None = Form(None),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Upload a document: validate, store in MinIO, create DB record, queue processing."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="Имя файла не указано")

    file_ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if file_ext not in ALLOWED_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise HTTPException(
            status_code=400,
            detail=f"Неподдерживаемый тип файла. Разрешены: {allowed}",
        )

    file_data = await file.read()
    if len(file_data) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="Файл превышает 50 MB")

    # Parse tags: accept JSON array or comma-separated string
    parsed_tags: list[str] = []
    if tags:
        try:
            parsed_tags = json.loads(tags)
            if not isinstance(parsed_tags, list):
                parsed_tags = []
        except (json.JSONDecodeError, ValueError):
            parsed_tags = [t.strip() for t in tags.split(",") if t.strip()]

    # Generate unique path in MinIO to avoid collisions
    minio_path = f"{uuid.uuid4()}/{file.filename}"
    content_type = CONTENT_TYPES.get(file_ext, "application/octet-stream")

    # Upload to MinIO in thread pool (synchronous boto3 call)
    try:
        await asyncio.get_running_loop().run_in_executor(
            None, upload_file, file_data, minio_path, content_type
        )
    except Exception as exc:
        logger.error("MinIO upload failed: %s", exc)
        raise HTTPException(status_code=500, detail="Ошибка загрузки файла в хранилище") from exc

    doc_title = title or file.filename
    document = Document(
        title=doc_title,
        filename=file.filename,
        file_type=file_ext,
        file_size=len(file_data),
        minio_path=minio_path,
        folder_path=folder_path,
        status="pending",
        uploaded_by=current_user.id,
        org_id=org_id,
        tags=parsed_tags,
        department=department,
    )
    db.add(document)
    await db.flush()  # assigns document.id

    # Uploader automatically gets owner permission
    perm = DocumentPermission(
        document_id=document.id,
        user_id=current_user.id,
        org_id=None,
        level="owner",
        granted_by=current_user.id,
    )
    db.add(perm)
    await db.commit()
    await db.refresh(document)

    # Queue Celery processing task
    _queue_process_task(document.id)

    return DocumentOut.model_validate(document)


@router.get("/", response_model=DocumentListOut)
async def list_documents(
    page: int = 1,
    limit: int = 20,
    doc_status: str | None = None,
    file_type: str | None = None,
    folder_path: str | None = None,
    org_id: int | None = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """List documents accessible to the current user with optional filters."""
    query = select(Document)

    if current_user.role != "admin":
        # Build access filter: own docs + org docs + explicit permissions
        org_result = await db.execute(
            select(OrganizationMember.org_id).where(OrganizationMember.user_id == current_user.id)
        )
        user_org_ids = [row[0] for row in org_result.all()]

        perm_result = await db.execute(
            select(DocumentPermission.document_id).where(
                DocumentPermission.user_id == current_user.id
            )
        )
        accessible_ids = [row[0] for row in perm_result.all()]

        if user_org_ids:
            org_perm_result = await db.execute(
                select(DocumentPermission.document_id).where(
                    DocumentPermission.org_id.in_(user_org_ids)
                )
            )
            accessible_ids += [row[0] for row in org_perm_result.all()]

        conditions = [Document.uploaded_by == current_user.id]
        if user_org_ids:
            conditions.append(Document.org_id.in_(user_org_ids))
        if accessible_ids:
            conditions.append(Document.id.in_(accessible_ids))

        query = query.where(or_(*conditions))

    if doc_status:
        query = query.where(Document.status == doc_status)
    if file_type:
        query = query.where(Document.file_type == file_type)
    if folder_path:
        query = query.where(Document.folder_path == folder_path)
    if org_id is not None:
        query = query.where(Document.org_id == org_id)

    count_result = await db.execute(select(func.count()).select_from(query.subquery()))
    total = count_result.scalar() or 0

    offset = (page - 1) * limit
    result = await db.execute(
        query.order_by(Document.uploaded_at.desc()).offset(offset).limit(limit)
    )
    documents = result.scalars().all()

    return DocumentListOut(
        items=[DocumentOut.model_validate(d) for d in documents],
        total=total,
        page=page,
        limit=limit,
    )


@router.get("/{document_id}", response_model=DocumentOut)
async def get_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get document metadata. Requires at least viewer access."""
    doc = await check_document_access(document_id, "viewer", current_user, db)
    return DocumentOut.model_validate(doc)


@router.get("/{document_id}/download")
async def download_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Download the raw file from MinIO. Requires viewer access."""
    doc = await check_document_access(document_id, "viewer", current_user, db)

    audit = AuditLog(
        user_id=current_user.id,
        action="download",
        resource_type="document",
        resource_id=str(document_id),
        details={"filename": doc.filename},
    )
    db.add(audit)

    try:
        file_bytes = await asyncio.get_running_loop().run_in_executor(
            None, download_file, doc.minio_path
        )
    except Exception as exc:
        logger.error("MinIO download failed for %s: %s", doc.minio_path, exc)
        raise HTTPException(status_code=500, detail="Ошибка получения файла из хранилища") from exc

    media_type = CONTENT_TYPES.get(doc.file_type, "application/octet-stream")
    return Response(
        content=file_bytes,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{doc.filename}"'},
    )


@router.get("/{document_id}/status", response_model=DocumentStatusOut)
async def get_document_status(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get processing status of a document."""
    doc = await check_document_access(document_id, "viewer", current_user, db)
    return DocumentStatusOut(
        id=doc.id,
        status=doc.status,
        error_message=doc.error_message,
        chunk_count=doc.chunk_count,
    )


@router.get("/{document_id}/preview", response_model=DocumentPreviewOut)
async def get_document_preview(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Return the first 3000 characters of extracted text. Requires viewer access."""
    doc = await check_document_access(document_id, "viewer", current_user, db)

    try:
        file_bytes = await asyncio.get_running_loop().run_in_executor(
            None, download_file, doc.minio_path
        )
        text = await asyncio.get_running_loop().run_in_executor(
            None, _extract_text_preview, file_bytes, doc.file_type
        )
    except Exception as exc:
        logger.error("Preview extraction failed for doc %d: %s", document_id, exc)
        raise HTTPException(status_code=500, detail="Ошибка извлечения текста") from exc

    return DocumentPreviewOut(text=text[:3000], page_count=doc.page_count)


@router.get("/{document_id}/presigned-url", response_model=PresignedUrlOut)
async def get_presigned_url(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Return a short-lived presigned URL for direct browser download/preview (viewer+).

    The URL is valid for 5 minutes and bypasses the API auth header requirement,
    making it usable in iframes and <a href> links.
    """
    doc = await check_document_access(document_id, "viewer", current_user, db)

    try:
        url = await asyncio.get_running_loop().run_in_executor(
            None, generate_presigned_url, doc.minio_path, 300
        )
    except Exception as exc:
        logger.error("Failed to generate presigned URL for doc %d: %s", document_id, exc)
        raise HTTPException(status_code=500, detail="Ошибка генерации URL") from exc

    return PresignedUrlOut(url=url, expires_in=300)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _extract_text_preview(file_data: bytes, file_type: str) -> str:
    """Extract text from a document for preview purposes."""
    if file_type == "pdf":
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_data, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)
    if file_type == "docx":
        import docx

        document = docx.Document(io.BytesIO(file_data))
        return "\n".join(p.text for p in document.paragraphs)
    if file_type == "xlsx":
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(file_data), read_only=True)
        lines = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                lines.append("\t".join("" if c is None else str(c) for c in row))
        return "\n".join(lines)
    # txt or unknown — decode as UTF-8
    return file_data.decode("utf-8", errors="replace")


def _queue_process_task(document_id: int) -> None:
    """Send the document processing task to the Celery broker."""
    try:
        from celery_client import celery_app

        celery_app.send_task(
            "worker.tasks.process_document.process_document",
            args=[document_id],
        )
        logger.info("Queued processing task for document %d", document_id)
    except Exception as exc:
        # Non-fatal: document stays in 'pending' state and can be reindexed later
        logger.warning("Failed to queue processing task for doc %d: %s", document_id, exc)
