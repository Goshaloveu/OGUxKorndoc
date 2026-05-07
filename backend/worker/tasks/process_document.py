"""
Celery task: full document processing pipeline.

Steps:
  1. Download file from MinIO to a temporary directory
  2. Set document status → "processing"
  3. Extract text based on file type (PyMuPDF / python-docx / openpyxl / plain text)
     with pytesseract OCR fallback for scan-only PDFs
  4. Sliding-window chunking (chunk_size / chunk_overlap from settings)
  5. Embed each chunk with the singleton sentence-transformer model
  6. Upsert chunks into Qdrant with access-control payload fields
  7. Update Document: status="indexed", indexed_at, chunk_count, page_count
  8. On error: status="error", error_message saved, retry after 60 s (max 3)
"""

from __future__ import annotations

import io
import logging
import uuid
from collections.abc import Generator
from datetime import UTC, datetime

from celery_app import celery_app, get_embedder, get_sparse_embedder
from shared.config import settings
from shared.minio_client import download_file

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Qdrant helpers (collection bootstrap + upsert)
# ---------------------------------------------------------------------------


def _get_qdrant_client():
    from qdrant_client import QdrantClient

    return QdrantClient(host=settings.qdrant_host, port=settings.qdrant_port)


def _ensure_collection(client) -> None:
    """Create the Qdrant collection if it does not exist yet."""
    from qdrant_client.http.exceptions import UnexpectedResponse

    existing = [c.name for c in client.get_collections().collections]
    if settings.qdrant_collection not in existing:
        try:
            _create_collection(client)
        except UnexpectedResponse as e:
            if e.status_code == 409:
                logger.info(
                    "Qdrant collection '%s' already exists (race condition ignored)",
                    settings.qdrant_collection,
                )
            else:
                raise
        return

    info = client.get_collection(settings.qdrant_collection)
    params = info.config.params
    vectors_config = params.vectors
    sparse_vectors_config = params.sparse_vectors or {}
    dense_config = vectors_config.get("dense") if isinstance(vectors_config, dict) else None
    sparse_config = (
        sparse_vectors_config.get("sparse") if isinstance(sparse_vectors_config, dict) else None
    )
    if dense_config is None or sparse_config is None:
        _handle_collection_mismatch(
            client,
            "must use named vectors 'dense' and 'sparse' for hybrid search",
        )
        return
    if dense_config.size != settings.embedding_dim:
        _handle_collection_mismatch(
            client,
            f"dense vector size is {dense_config.size}, expected {settings.embedding_dim}",
        )


def _create_collection(client) -> None:
    from qdrant_client.http.models import (
        Distance,
        SparseIndexParams,
        SparseVectorParams,
        VectorParams,
    )

    client.create_collection(
        collection_name=settings.qdrant_collection,
        vectors_config={
            "dense": VectorParams(size=settings.embedding_dim, distance=Distance.COSINE)
        },
        sparse_vectors_config={
            "sparse": SparseVectorParams(index=SparseIndexParams(on_disk=False))
        },
    )
    logger.info(
        "Created Qdrant collection '%s' (dense dim=%d, sparse=%s)",
        settings.qdrant_collection,
        settings.embedding_dim,
        settings.sparse_embedding_model,
    )


def _handle_collection_mismatch(client, reason: str) -> None:
    message = f"Qdrant collection '{settings.qdrant_collection}' {reason}."
    if not settings.allow_qdrant_recreate_on_mismatch:
        logger.error("%s Recreate it and run reindex_all_documents.", message)
        raise RuntimeError(f"{message} Recreate it and run reindex_all_documents.")

    logger.warning(
        "%s Recreating it because ENVIRONMENT=%s. Run reindex_all_documents after startup.",
        message,
        settings.environment,
    )
    client.delete_collection(collection_name=settings.qdrant_collection)
    _create_collection(client)


def _recreate_collection(client) -> None:
    """Explicitly recreate the collection before a full migration reindex."""
    existing = [c.name for c in client.get_collections().collections]
    if settings.qdrant_collection in existing:
        client.delete_collection(collection_name=settings.qdrant_collection)

    _create_collection(client)
    logger.info("Recreated Qdrant collection '%s' for hybrid search", settings.qdrant_collection)


def _delete_document_points(client, document_id: int) -> None:
    """Remove existing chunks for a document before reindexing it."""
    from qdrant_client.http.models import FieldCondition, Filter, MatchValue

    client.delete(
        collection_name=settings.qdrant_collection,
        points_selector=Filter(
            must=[FieldCondition(key="document_id", match=MatchValue(value=document_id))]
        ),
    )


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


def _extract_text_pdf(file_data: bytes) -> tuple[str, int, list[str]]:
    """Extract text from PDF.  Falls back to OCR if the page has no text layer."""
    import fitz  # PyMuPDF

    doc = fitz.open(stream=file_data, filetype="pdf")
    pages: list[str] = []
    for page in doc:
        text = page.get_text().strip()
        if not text:
            # OCR fallback via pytesseract
            try:
                import pytesseract
                from PIL import Image

                pix = page.get_pixmap(dpi=150)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text = pytesseract.image_to_string(img, lang="rus+eng")
            except Exception as ocr_err:
                logger.warning("OCR failed for page %d: %s", page.number, ocr_err)
                text = ""
        pages.append(text)
    return "\n\n".join(pages), len(pages), pages


def _extract_text_docx(file_data: bytes) -> tuple[str, int, list[str]]:
    import docx

    document = docx.Document(io.BytesIO(file_data))

    parts: list[str] = []

    # Extract text from paragraphs
    for p in document.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # Extract text from tables
    for table in document.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append("\t".join(row_texts))

    text = "\n".join(parts)
    logger.info(
        "DOCX extraction: %d paragraphs, %d tables, %d chars total",
        len(document.paragraphs),
        len(document.tables),
        len(text),
    )

    # Fallback: if still empty, try raw XML extraction via zipfile
    if not text.strip():
        logger.warning("python-docx gave empty text, falling back to raw XML extraction")
        import re
        import zipfile

        try:
            with zipfile.ZipFile(io.BytesIO(file_data)) as zf:
                if "word/document.xml" in zf.namelist():
                    xml_content = zf.read("word/document.xml").decode("utf-8", errors="replace")
                    # Extract text between <w:t> tags
                    raw_texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml_content, re.DOTALL)
                    text = " ".join(raw_texts)
                    logger.info("XML fallback extracted %d chars", len(text))
        except Exception as xml_err:
            logger.warning("XML fallback also failed: %s", xml_err)

    return text, 0, [text]


def _extract_text_xlsx(file_data: bytes) -> tuple[str, int, list[str]]:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(file_data), read_only=True)
    lines: list[str] = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            lines.append("\t".join("" if c is None else str(c) for c in row))
    text = "\n".join(lines)
    return text, 0, [text]


def _extract_text(file_data: bytes, file_type: str) -> tuple[str, int, list[str]]:
    """Return (full_text, page_count).  page_count is 0 for non-PDF formats."""
    if file_type == "pdf":
        return _extract_text_pdf(file_data)
    if file_type == "docx":
        return _extract_text_docx(file_data)
    if file_type == "xlsx":
        return _extract_text_xlsx(file_data)
    # txt / fallback
    text = file_data.decode("utf-8", errors="replace")
    return text, 0, [text]


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------


def _chunk_text(
    text: str,
    pages: list[str],
    chunk_size: int = 2048,
    overlap: int = 256,
) -> Generator[tuple[str, int], None, None]:
    """Yield sliding-window chunks of *text* with their source page number."""
    page_offsets: list[tuple[int, int]] = []
    offset = 0
    for page_index, page_text in enumerate(pages, start=1):
        page_offsets.append((offset, page_index))
        offset += len(page_text) + 2

    start = 0
    text_len = len(text)
    while start < text_len:
        end = min(start + chunk_size, text_len)
        page = 0
        if len(pages) > 1:
            for page_offset, page_index in page_offsets:
                if page_offset <= start:
                    page = page_index
                else:
                    break
        yield text[start:end], page
        if end == text_len:
            break
        start += chunk_size - overlap


def _to_qdrant_sparse_vector(sparse_embedding):
    from qdrant_client.http.models import SparseVector

    return SparseVector(
        indices=sparse_embedding.indices.tolist(),
        values=sparse_embedding.values.tolist(),
    )


# ---------------------------------------------------------------------------
# Synchronous DB helpers (called from Celery task, not async context)
# ---------------------------------------------------------------------------


def _update_document_status(
    document_id: int,
    status: str,
    *,
    error_message: str | None = None,
    chunk_count: int | None = None,
    page_count: int | None = None,
    indexed_at: datetime | None = None,
) -> None:
    """Synchronously update document fields using a psycopg2 connection."""
    import psycopg2

    dsn = (
        f"host={settings.postgres_host} port={settings.postgres_port} "
        f"dbname={settings.postgres_db} user={settings.postgres_user} "
        f"password={settings.postgres_password}"
    )
    conn = psycopg2.connect(dsn)
    try:
        with conn.cursor() as cur:
            fields = ["status = %s", "updated_at = now()"]
            values: list = [status]

            if error_message is not None:
                fields.append("error_message = %s")
                values.append(error_message)
            if chunk_count is not None:
                fields.append("chunk_count = %s")
                values.append(chunk_count)
            if page_count is not None:
                fields.append("page_count = %s")
                values.append(page_count)
            if indexed_at is not None:
                fields.append("indexed_at = %s")
                values.append(indexed_at)

            values.append(document_id)
            cur.execute(
                f"UPDATE documents SET {', '.join(fields)} WHERE id = %s",  # noqa: S608
                values,
            )
        conn.commit()
    finally:
        conn.close()


def _get_document_info(document_id: int) -> dict | None:
    """Fetch the minimal document fields needed for processing."""
    import psycopg2
    import psycopg2.extras

    dsn = (
        f"host={settings.postgres_host} port={settings.postgres_port} "
        f"dbname={settings.postgres_db} user={settings.postgres_user} "
        f"password={settings.postgres_password}"
    )
    conn = psycopg2.connect(dsn)
    try:
        with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(
                "SELECT id, title, minio_path, file_type, folder_path, uploaded_by, org_id "
                "FROM documents WHERE id = %s",
                (document_id,),
            )
            row = cur.fetchone()
            return dict(row) if row else None
    finally:
        conn.close()


# ---------------------------------------------------------------------------
# Celery task
# ---------------------------------------------------------------------------


@celery_app.task(
    bind=True,
    name="worker.tasks.process_document.process_document",
    max_retries=3,
    default_retry_delay=60,
)
def process_document(self, document_id: int) -> None:
    """
    Full document processing pipeline:
    download → extract text → chunk → embed → upsert Qdrant → update status.
    """
    logger.info("Starting processing for document_id=%d", document_id)

    # 1. Fetch document metadata from PostgreSQL
    doc_info = _get_document_info(document_id)
    if doc_info is None:
        logger.error("Document %d not found in DB — skipping", document_id)
        return

    minio_path: str = doc_info["minio_path"]
    file_type: str = doc_info["file_type"]
    title: str = doc_info["title"]
    folder_path: str = doc_info["folder_path"]
    uploaded_by: int = doc_info["uploaded_by"]
    org_id: int | None = doc_info["org_id"]

    # 2. Update status → "processing"
    _update_document_status(document_id, "processing")

    try:
        # 3. Download from MinIO
        logger.info("Downloading %s from MinIO", minio_path)
        file_data = download_file(minio_path)

        # 4. Extract text
        logger.info("Extracting text from %s (type=%s)", minio_path, file_type)
        full_text, page_count, pages = _extract_text(file_data, file_type)

        if not full_text.strip():
            logger.warning("Document %d produced empty text after extraction", document_id)

        # 5. Chunking
        chunks = list(_chunk_text(full_text, pages, settings.chunk_size, settings.chunk_overlap))
        logger.info("Document %d split into %d chunks", document_id, len(chunks))

        qdrant = _get_qdrant_client()
        _ensure_collection(qdrant)
        _delete_document_points(qdrant, document_id)

        if not chunks:
            _update_document_status(
                document_id,
                "indexed",
                chunk_count=0,
                page_count=page_count if page_count else None,
                indexed_at=datetime.now(tz=UTC),
            )
            logger.info("Document %d indexed with no searchable chunks", document_id)
            return

        # 6. Embed all chunks at once (batch for efficiency)
        embedder = get_embedder()
        chunk_texts = [chunk_text for chunk_text, _page in chunks]
        vectors = embedder.encode(chunk_texts, show_progress_bar=False, batch_size=32)
        sparse_embedder = get_sparse_embedder()
        sparse_vectors = list(sparse_embedder.embed(chunk_texts))

        # 7. Upsert into Qdrant
        from qdrant_client.http.models import PointStruct

        points: list[PointStruct] = []
        for idx, ((chunk_text, page), vector, sparse_vector) in enumerate(
            zip(chunks, vectors, sparse_vectors, strict=True)
        ):
            point_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{document_id}-{idx}"))
            points.append(
                PointStruct(
                    id=point_id,
                    vector={
                        "dense": vector.tolist(),
                        "sparse": _to_qdrant_sparse_vector(sparse_vector),
                    },
                    payload={
                        "document_id": document_id,
                        "title": title,
                        "chunk_index": idx,
                        "page": page,
                        "text": chunk_text,
                        "file_type": file_type,
                        "folder_path": folder_path,
                        "uploaded_by": uploaded_by,
                        "org_id": org_id,
                    },
                )
            )

        # Upsert in batches of 100 to avoid large payloads
        batch_size = 100
        for i in range(0, len(points), batch_size):
            qdrant.upsert(
                collection_name=settings.qdrant_collection,
                points=points[i : i + batch_size],
            )
        logger.info("Upserted %d vectors to Qdrant for document %d", len(points), document_id)

        # 8. Update document: indexed
        _update_document_status(
            document_id,
            "indexed",
            chunk_count=len(chunks),
            page_count=page_count if page_count else None,
            indexed_at=datetime.now(tz=UTC),
        )
        logger.info("Document %d successfully indexed (%d chunks)", document_id, len(chunks))

    except Exception as exc:
        logger.exception("Error processing document %d: %s", document_id, exc)
        _update_document_status(
            document_id,
            "error",
            error_message=str(exc)[:1000],
        )
        raise self.retry(exc=exc, countdown=60) from None


@celery_app.task(
    name="worker.tasks.process_document.reindex_all_documents",
)
def reindex_all_documents() -> dict[str, int]:
    """Queue all non-deleted documents for reindexing after collection/vector schema changes."""
    import psycopg2

    qdrant = _get_qdrant_client()
    _recreate_collection(qdrant)

    dsn = (
        f"host={settings.postgres_host} port={settings.postgres_port} "
        f"dbname={settings.postgres_db} user={settings.postgres_user} "
        f"password={settings.postgres_password}"
    )
    conn = psycopg2.connect(dsn)
    try:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE documents "
                "SET status = 'pending', error_message = NULL, chunk_count = NULL, "
                "indexed_at = NULL, updated_at = now() "
                "WHERE status != 'deleted' "
                "RETURNING id"
            )
            document_ids = [row[0] for row in cur.fetchall()]
        conn.commit()
    finally:
        conn.close()

    for document_id in document_ids:
        process_document.delay(document_id)

    logger.info("Queued %d documents for hybrid search reindexing", len(document_ids))
    return {"queued": len(document_ids)}
